# routes/gemai.py
import os
from datetime import datetime
from typing import Optional, List
from fastapi import APIRouter
from pydantic import BaseModel
import re
from difflib import get_close_matches
from datetime import datetime

from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
import docx

from database.create_tables import SessionLocal
from database.models import GemAILog


from openai import OpenAI
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

router = APIRouter()

# ---------------------- Utility Functions ----------------------

# Utility functions
def load_portfolio_docs() -> List[Document]:
   docs = []
   folder_path = "./portfolio_embeds"
   for filename in os.listdir(folder_path):
       if filename.endswith(".docx"):
           doc = docx.Document(os.path.join(folder_path, filename))
           full_text = "\n".join([p.text for p in doc.paragraphs])
           docs.append(Document(page_content=full_text, metadata={"source": filename}))
   return docs


def extract_latest_project(docs: List[Document]) -> Optional[str]:
   month_map = {
       'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
       'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
   }


   date_project_pairs = []

   for doc in docs:
       if "resume" in doc.metadata["source"].lower():
           lines = doc.page_content.splitlines()
           for i in range(len(lines) - 1):
               title = lines[i].strip()
               date_line = lines[i + 1].strip()


               # ❌ Skip experience or job/company headers
               if title.upper() in {
                   "LEADERSHIP EXPERIENCE", "PROFESSIONAL EXPERIENCE",
                   "INTERNSHIP EXPERIENCE", "WORK EXPERIENCE",
                   "EXPERIENCE", "ACADEMIC PROJECTS / PERSONAL PROJECTS"
               } or re.search(r"(Infosys|Company|Kathmandu|Engineer|Developer|Intern)", title, re.IGNORECASE):
                   continue


               # ✅ Match project date format
               match = re.search(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b\s+(\d{4})", date_line, re.IGNORECASE)
               if match and len(title.split()) <= 10:
                   try:
                       month = match.group(1).lower()
                       year = int(match.group(2))
                       dt = datetime(year, month_map[month], 1)
                       date_project_pairs.append((dt, title, f"{match.group(1)} {year}"))
                   except:
                       continue


   if date_project_pairs:
       latest = sorted(date_project_pairs, reverse=True)[0]
       return f"Most recent project:\n\n{latest[1]} ({latest[2]})"


   return None


def extract_text_from_docx(docx_path: str) -> str:
   doc = docx.Document(docx_path)
   return "\n".join([p.text for p in doc.paragraphs])


# ---------------------- Vector Store Setup ----------------------


def setup_vectorstore() -> FAISS:
   docs = load_portfolio_docs()
  # Load both resume files
   docx_text = extract_text_from_docx("./public/PrashantResume.docx")


   docs.append(Document(page_content=docx_text, metadata={"source": "Resume DOCX"}))


   # Chunk and embed
   splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
   split_docs = splitter.split_documents(docs)
   embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
   return FAISS.from_documents(split_docs, embeddings)


vectorstore = setup_vectorstore()


# ---------------------- Global Portfolio Docs ----------------------
portfolio_docs = load_portfolio_docs()


# ---------------------- AI Agent Endpoint ----------------------
class QueryRequest(BaseModel):
   query: str
   session_id: Optional[str] = None 


# Small talk handler
def handle_small_talk(query: str) -> Optional[str]:
   q = query.lower().strip()

   # Remove punctuation for flexible matching
   q_clean = re.sub(r"[^\w\s]", "", q)

   # Greetings
   if q_clean.strip() in {"hi", "hello", "hey", "whats up", "what's up", "yo", "hiya"}:
       return (
           "Hello! I am GemAI, here to help you explore Prashant's portfolio.\n\n"
           "For example:\n"
           "• what are Prashant's AI/ML projects?\n"
           "• Prashant's resume\n"
           "• How to contact Prashant?"
       )

   # Thanks
   if "thank" in q_clean:
       return "You're welcome! 😊 Let me know if you'd like help with Prashant's resume, skills, or contact info."

   # How are you
   if "how are you" in q_clean or "how you doing" in q_clean:
       return "I am doing great! Let me know how I can assist with Prashant's profile."

   # AI-related identity questions
   if any(phrase in q_clean for phrase in [
       "who are you", "what is your name",
       "are you ai", "are you an ai", "are you real", "are you human"
   ]):
       return "I am GemAI — the AI agent embedded in Prashant Baral's portfolio. Ask me anything about his skills, resume, or work."

   # Apologies
   if "sorry" in q_clean:
       return "No worries at all! Let me know if you need help exploring Prashant's profile."

   # Sensitive topics
   if "personal life" in q_clean:
       return "Sorry, I couldn not find personal life details in Prashant's profile. But I can share his hobbies and interests if you'd like!"

   return None

def match_keywords(query: str, keywords: dict, aliases: dict) -> List[str]:
   query_lower = query.lower()
   words = re.findall(r'\b\w+\b', query_lower)
   matched = []
   for key in keywords:
       if key in query_lower or get_close_matches(key, words, n=1, cutoff=0.8):
           matched.append(keywords[key])
   for alias in aliases:
       if alias in query_lower or get_close_matches(alias, words, n=1, cutoff=0.8):
           if aliases[alias] in keywords:
               matched.append(keywords[aliases[alias]])
   return matched


def tech_in_profile(tech: str, docs: List[Document]) -> bool:
   """
   Check if the given technology keyword is explicitly mentioned in the loaded portfolio documents.
   Supports case-insensitive exact matches.
   """
   pattern = re.compile(rf"\b{re.escape(tech)}\b", re.IGNORECASE)
   for doc in docs:
       if pattern.search(doc.page_content):
           return True
   return False


@router.post("/ai-agent/")
async def ai_agent(request: QueryRequest):
   query = request.query.strip()
   session_id = request.session_id  # <- capture it

   # Normalize common typo variants
   qnorm = query.lower().strip()

   small_talk_response = handle_small_talk(query)
   if small_talk_response:
       return {"response": small_talk_response}
  

   # Block forbidden topics
   forbidden_words = ["religion", "politics", "violence", "hate", "intimacy", "fight", "abuse", "income", "dating", "boyfriend", "age"]
   if any(word in qnorm for word in forbidden_words):
       return {"response": "This topic is not part of the portfolio."}


   # Step 1: Retrieve top relevant content
   keywords = {
       "services": "services.docx",
       "skills": "skills.docx",
       "projects": "projects.docx",
       "interest": "interests.docx",
       "resume": "resume.docx",
       "education": "qualification.docx",
       "work": "resume.docx",
       "qualification": "qualification.docx",
   }

   aliases = {
       "service": "services",
       "services": "services",
       "offer": "services",
       "help": "services",
       "give": "services",
       "skill": "skills",
       "project": "projects",
       "qualification": "qualifications",
       "experience": "work",
   }
  
   matched_keywords = match_keywords(qnorm, keywords, aliases)
  
   # Handle vague identity queries and bios by forcing resume + interests context
   vague_patterns = [
       r"(detail|details|about|info|information)\s+(of|on|regarding)?\s*(Prashant|his|Prashant Baral)",
       r"(describe|describes|tells|tell)\s+(about\s+)?(Prashant|his)",
       r"(who\s+(is|are)\s+)?Prashant( Baral)?", r"who is he", r"his bio"
   ]
   if any(re.search(p, qnorm) for p in vague_patterns):
       matched_keywords += ["qualification.docx", "interests.docx"]
  
   extra_context = []


   # Add matched keyword files (e.g., skills.docx, projects.docx)
   for doc in portfolio_docs:
       if doc.metadata["source"] in matched_keywords:
            extra_context.append(doc)


   # Ensure skills.docx is always included (used for tech stack confirmation)
   if not any(doc.metadata["source"] == "skills.docx" for doc in extra_context):
       for doc in portfolio_docs:
           if doc.metadata["source"] == "skills.docx":
                extra_context.append(doc)
                break



   # Step 2: Vector search
   results = vectorstore.similarity_search(qnorm, k=3)
   combined_docs = results + extra_context
   if not combined_docs:
       return {"response": "I couldn't find information related to that in Prashant's profile."}


   # Handle queries specifically asking about latest or most recent project
   if re.search(r"(latest|recent|worked on|academic|newest|project he|he worked)", qnorm) and "project" in qnorm:
       latest_proj = extract_latest_project(combined_docs)
       if latest_proj:
           return {"response": latest_proj}


  
   context = "\n\n".join([doc.page_content for doc in combined_docs])


   # Step 3: Construct better prompt with markdown link instructions and hallucination warning
   prompt = f"""
You are **GemAI**, the embedded AI agent for Prashant Baral's portfolio website. You must only use the information provided in the context below to answer the user's question.




📌 **If the answer is not found in the context, DO NOT guess or make assumptions. Respond with:**
"I could not find information related to that in Prashant's profile."


📌 **You are GemAI, Prashant's portfolio AI agent. NEVER say "I am Prashant Baral". NEVER answer in first person as if you are him. Always refer to his as "Prashant Baral" or "he".**


📌 **When giving links, always format using markdown**  and **only mention once if necessary**:
- [Resume (PDF)](/files/PrashantResume.pdf)
- [Personal Email](mailto:thedashh7@gmail.com)
- [University Email](mailto:pb23395n@pace.edu)
- [LinkedIn](https://www.linkedin.com/in/prashant-baral/)
- [GitHub](https://github.com/baralprashant)



📌 **Never mention external platforms like Instagram, Spotify, Meetup, TripAdvisor, or any websites unless they are explicitly present in the context.**


📌 **Do not include links like "Click here", "Learn more", or vague demo or portfolio pages unless they exist in the provided context.**


📌 **Never mention sections like “Interests & Hobbies” or “portfolio website links” unless the exact phrase is found in the text.**


📌 **Never hyperlink the text "Prashant Baral". Only show his name in plain text.**


📌 If the user asks vague queries like “Tell me about him”, “Who is Prashant”, “Describe Prashant”, “Details of Prashant Baral”, or “About Prashant Baral”, then :**
- **Write a SINGLE short paragraph, no more than 3 to 4 concise lines.**
- combine details from :
- `qualification.docs` + resume for education/work
- `interests.docs` for personal interests
- Keep it brief, clear, and professional
- End with GitHub, LinkedIn, and resume links


📌 If asked about “qualification” then give content from only `qualification.docx`.
- Include both educational background and work experience from `qualification.docx` and nothing more than that.
- Summarize naturally in a clear format.


📌 If asked about “latest”, “recent”, or “academic” project, return only one academic/personal project from resume.docx or projects.docx.**
- Never confuse it with “LEADERSHIP” or “PROFESSIONAL EXPERIENCE”.
- Use project title and month/year only
📌 Special rule: Ignore company or experience headers like “Javra Software Pvt. Ltd.” when user asks about “latest project” or “academic project”. Only return a real project title with a month and year.


📌 **For tech skills, check both files skills.docx and resume.docx . If found, answer naturally — you can summarize what he used it for or how it is described. If not found, say you could not find it. Never guess.
- Do not mention it is listed.

📌 **Keep responses structured, neutral, and based ONLY on the provided `.docx` files and resume content.** Avoid fabricating details or repeating contact links unless specifically asked.


- "resume": provide a clickable markdown resume download link
- "skills": summarize only the `skills.docx` content
- "projects": summarize only `projects.docx` and resume projects
- "education": from `qualification.docx` or resume.docx
- "work": summarize from `qualification.docx` and resume.docx
- “services” → provide what's in the services.docx
- "qualification": summarize only from qualification.docx
- “interests”, “hobbies”, “free time” → summarize interests.docx only once




Context:
{context}


User Question:
{qnorm}


Answer using markdown only. Never invent or guess information not in the context.
"""


   try:
       response = client.chat.completions.create(
           model="gpt-4", # or "gpt-3.5-turbo",
           messages=[ {"role": "system", "content": "You are GemAI, the embedded AI agent for Prashant Baral’s portfolio website."}, {"role": "user", "content": prompt} ],
           temperature=0.2,
           max_tokens=700
       )
       answer = response.choices[0].message.content


       # Fix: Strip hyperlinks from Prashant’s name
       answer = re.sub(r"\[Prashant Baral\]\([^)]+\)", "Prashant Baral", answer)


   except Exception as e:
       answer = f" Error: {str(e)}"


   session = SessionLocal()
   log = GemAILog(
       session_id=session_id,  # <- store it
       timestamp=datetime.utcnow().isoformat(),
       question=query,
       response=answer,
      
   )
   session.add(log)
   session.commit()
   session.close()


   return {"response": answer}